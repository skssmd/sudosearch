import os
import sys
import tkinter as tk
from tkinter import filedialog
from fuzzywuzzy import fuzz
from docx import Document
from PyPDF2 import PdfReader
import pytesseract
from PIL import Image
from pdf2image import convert_from_path
import openpyxl
import subprocess

def select_folder():
    root = tk.Tk()
    root.withdraw()
    folder = filedialog.askdirectory(title="Select folder to search")
    return folder

def get_file_text(file_path):
    try:
        if file_path.endswith(".txt"):
            with open(file_path, "r", errors="ignore") as f:
                return f.read()
        elif file_path.endswith(".docx"):
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            # Extract table text
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    text += "\n" + row_text
            return text

        elif file_path.endswith(".pdf"):
            import pdfplumber
            text = ""
            try:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except:
                return ""
            return text
        elif file_path.endswith(".xlsx"):
            wb = openpyxl.load_workbook(file_path, data_only=True)
            text = ""
            for sheet in wb.worksheets:
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    text += row_text + "\n"
            return text
    except Exception as e:
        return ""
    return ""

def ocr_image_or_pdf(file_path):
    try:
        if file_path.lower().endswith((".png", ".jpg", ".jpeg", ".bmp")):
            return pytesseract.image_to_string(Image.open(file_path))
        elif file_path.lower().endswith(".pdf"):
            text = ""
            pages = convert_from_path(file_path)
            for img in pages:
                text += pytesseract.image_to_string(img)
            return text
    except:
        return ""
    return ""

def fuzzy_search(text, keyword):
    return fuzz.partial_ratio(text.lower(), keyword.lower()) >= 80

import time
import threading

def search_files(base_dir, search_type, keyword):
    matched = []
    all_files = []

    # Collect all files first
    for root, dirs, files in os.walk(base_dir):
        for file in files:
            all_files.append(os.path.join(root, file))

    total = len(all_files)
    searched = [0]
    found = [0]

    def print_status():
        while searched[0] < total:
            print(f"\rTotal: {total} | Searched: {searched[0]} | Found: {found[0]} | Remaining: {total - searched[0]}", end="")
            time.sleep(0.2)
        print(f"\rTotal: {total} | Searched: {searched[0]} | Found: {found[0]} | Remaining: {total - searched[0]}", end="")

    status_thread = threading.Thread(target=print_status)
    status_thread.start()

    for full_path in all_files:
        try:
            matched_flag = False
            filename = os.path.basename(full_path)

            if search_type == "filename":
                if fuzzy_search(filename, keyword):
                    matched.append(full_path)
                    matched_flag = True
            elif search_type == "text":
                if filename.lower().endswith((".txt", ".docx", ".pdf", ".xlsx")):
                    content = get_file_text(full_path)
                    print(f"\nChecking file: {filename}")
                    print("[FOUND TEXT]")
                    print(content[:500] + ("..." if len(content) > 500 else ""))
                    if fuzzy_search(content, keyword):
                        matched.append((full_path, content))
                        matched_flag = True
                        print("Found match in this file!")

            elif search_type == "scanned":
                if filename.lower().endswith((".pdf", ".png", ".jpg", ".jpeg", ".bmp")):
                    content = ocr_image_or_pdf(full_path)
                    if fuzzy_search(content, keyword):
                        matched.append(full_path)
                        matched_flag = True
            elif search_type == "all":
                if fuzzy_search(filename, keyword):
                    matched.append(full_path)
                    matched_flag = True
                elif filename.lower().endswith((".txt", ".docx", ".pdf", ".xlsx")):
                    content = get_file_text(full_path)
                    if fuzzy_search(content, keyword):
                        matched.append(full_path)
                        matched_flag = True
                elif filename.lower().endswith((".pdf", ".png", ".jpg", ".jpeg", ".bmp")):
                    content = ocr_image_or_pdf(full_path)
                    if fuzzy_search(content, keyword):
                        matched.append(full_path)
                        matched_flag = True
        except:
            pass

        searched[0] += 1
        if matched_flag:
            found[0] += 1

    status_thread.join()
    print()

    # Dump all texts of matched files to a file
    with open("sudosearch_results.txt", "w", encoding="utf-8") as f:
        for path, content in matched:
            f.write(f"===== FILE: {path} =====\n")
            f.write(content + "\n\n")

    return [path for path, _ in matched]  # Return only file paths

    return matched


def open_options(filepath):
    while True:
        choice = input("1) Open File\n2) Open Location\nChoose (1/2): ")
        if choice == "1":
            os.startfile(filepath)
            break
        elif choice == "2":
            subprocess.run(f'explorer /select,"{filepath}"')
            break
        else:
            print("Invalid choice.")

def run_cli():
    print("=== SudoSearch ===")
    folder = select_folder()
    if not folder:
        print("No folder selected. Exiting.")
        return

    print("Search from:\n1) Filename\n2) Text\n3) Scanned\n4) All")
    type_map = {"1": "filename", "2": "text", "3": "scanned", "4": "all"}
    choice = input("Enter your choice (1-4): ").strip()

    if choice not in type_map:
        print("Invalid choice.")
        return

    search_type = type_map[choice]
    print(f"Searching in: {folder} for {search_type}")
    keyword = input("Please enter keyword to search: ").strip()

    results = search_files(folder, search_type, keyword)

    if not results:
        print("No matches found.")
        return

    print("\nSearch Results:")
    for i, path in enumerate(results, 1):
        print(f"{i}) {path}")

    try:
        sel = int(input("Select one to open (number): "))
        if 1 <= sel <= len(results):
            open_options(results[sel - 1])
        else:
            print("Invalid selection.")
    except:
        print("Invalid input.")

if __name__ == "__main__":
    if len(sys.argv) == 1:
        run_cli()
    else:
        # Reserved for future use like `sudosearch location type keyword`
        print("Direct CLI args not yet implemented. Use: `sudosearch` only for now.")
